import requests
import tempfile
from fpdf import FPDF

OLLAMA_URL = "http://localhost:11434/api/generate"

# ─────────────────────────────────────────────────────────────────────────────
# Genre Recommendation
# ─────────────────────────────────────────────────────────────────────────────
def recommend_genre(idea: str) -> str:
    prompt = f"""You are a creative writing expert and story analyst.

A user has the following story idea:
\"\"\"{idea}\"\"\"

Based on this idea, suggest 3–5 suitable story genres. For each genre:
- State the genre name
- Give a one-sentence explanation of why it fits the idea
- Rate how well it fits (e.g., Perfect fit / Good fit / Possible fit)

Format the output as a clean markdown list. Be concise and insightful.
"""
    return ask_llama(prompt, temperature=0.5, max_tokens=300)



# ─────────────────────────────────────────────────────────────────────────────
# LLM
# ─────────────────────────────────────────────────────────────────────────────

def ask_llama(prompt, temperature=0.8, max_tokens=None):
    data = {
        "model": "llama3",
        "prompt": prompt,
        "stream": False,
        "options": {
            "temperature": temperature,
            "num_ctx": 8192
        }
    }
    if max_tokens:
        data["options"]["num_predict"] = max_tokens

    try:
        r = requests.post(OLLAMA_URL, json=data, timeout=120)
        r.raise_for_status()
        return r.json()["response"]
    except requests.exceptions.ConnectionError:
        return "ERROR: ⚠ AI model not available. Please start Ollama."
    except requests.exceptions.Timeout:
        return "ERROR: ⏱ Request timed out. The model may be overloaded."
    except Exception as e:
        return f"ERROR: An unexpected error occurred: {str(e)}"


# ─────────────────────────────────────────────────────────────────────────────
# Export helpers
# ─────────────────────────────────────────────────────────────────────────────

def _meta_header(genre, characters, content_type):
    """Return a list of (label, value) metadata rows."""
    from datetime import date
    return [
        ("Format",     content_type or "Story"),
        ("Genre",      genre or "—"),
        ("Characters", characters or "—"),
        ("Date",       date.today().strftime("%d %B %Y")),
    ]


def create_pdf(text, genre="", characters="", content_type=""):
    safe = lambda s: s.encode("latin-1", "replace").decode("latin-1")
    meta = _meta_header(genre, characters, content_type)

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.add_page()

    # Title
    pdf.set_font("Arial", "B", 20)
    pdf.cell(0, 12, safe("AI STORY GENERATOR"), ln=True, align="C")
    pdf.set_draw_color(80, 80, 80)
    pdf.line(10, pdf.get_y(), 200, pdf.get_y())
    pdf.ln(4)

    # Metadata rows
    for label, value in meta:
        pdf.set_font("Arial", "B", 11)
        pdf.cell(38, 7, safe(f"{label}:"), border=0)
        pdf.set_font("Arial", "", 11)
        pdf.cell(0, 7, safe(value), ln=True)

    pdf.ln(4)
    pdf.line(10, pdf.get_y(), 200, pdf.get_y())
    pdf.ln(6)

    # Body
    pdf.set_font("Arial", "", 12)
    for line in text.split("\n"):
        pdf.multi_cell(0, 8, safe(line))

    # Page number footer
    pdf.set_font("Arial", "I", 9)
    pdf.set_y(-15)
    pdf.cell(0, 10, safe(f"Page {pdf.page_no()}"), align="C")

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")
    pdf.output(tmp.name)
    return tmp.name


def create_docx(text, genre="", characters="", content_type=""):
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    meta = _meta_header(genre, characters, content_type)

    doc = Document()

    # Title
    heading = doc.add_heading("AI STORY GENERATOR", 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Metadata table
    tbl = doc.add_table(rows=len(meta), cols=2)
    tbl.style = "Table Grid"
    for i, (label, value) in enumerate(meta):
        tbl.rows[i].cells[0].text = label
        tbl.rows[i].cells[1].text = value

    doc.add_paragraph("")  # spacer

    # Body
    for line in text.split("\n"):
        doc.add_paragraph(line)

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(tmp.name)
    return tmp.name


def create_txt(text, genre="", characters="", content_type=""):
    meta = _meta_header(genre, characters, content_type)
    lines = ["AI STORY GENERATOR", "=" * 40]
    for label, value in meta:
        lines.append(f"{label:<12}: {value}")
    lines += ["=" * 40, "", text]

    tmp = tempfile.NamedTemporaryFile(
        delete=False, suffix=".txt", mode="w", encoding="utf-8"
    )
    tmp.write("\n".join(lines))
    tmp.flush()
    return tmp.name


# ─────────────────────────────────────────────────────────────────────────────
# Voice Narration (gTTS)
# ─────────────────────────────────────────────────────────────────────────────

def text_to_audio(text: str, lang_code: str = "en") -> bytes:
    """Convert text to MP3 audio bytes using gTTS. Returns bytes or raises."""
    from gtts import gTTS
    import io
    tts = gTTS(text=text[:3000], lang=lang_code, slow=False)   # cap at 3000 chars
    buf = io.BytesIO()
    tts.write_to_fp(buf)
    buf.seek(0)
    return buf.read()


GTTS_LANG_CODES = {
    "English":  "en",
    "Tamil":    "ta",
    "Hindi":    "hi",
}


# ─────────────────────────────────────────────────────────────────────────────
# SQLite Story Library
# ─────────────────────────────────────────────────────────────────────────────

import sqlite3, os, pathlib

DB_PATH = pathlib.Path(__file__).parent / "story_library.db"

def _get_conn():
    conn = sqlite3.connect(str(DB_PATH))
    conn.row_factory = sqlite3.Row
    return conn

def init_db():
    """Create the stories table if it does not yet exist."""
    with _get_conn() as conn:
        conn.execute("""
            CREATE TABLE IF NOT EXISTS stories (
                id          INTEGER PRIMARY KEY AUTOINCREMENT,
                title       TEXT,
                genre       TEXT,
                characters  TEXT,
                content_type TEXT,
                language    TEXT,
                text        TEXT NOT NULL,
                saved_at    TEXT DEFAULT (datetime('now','localtime'))
            )
        """)
        conn.commit()

def save_story(title: str, text: str, genre="", characters="",
               content_type="", language="English") -> int:
    """Insert a story and return its new row id."""
    init_db()
    with _get_conn() as conn:
        cur = conn.execute(
            "INSERT INTO stories (title, genre, characters, content_type, language, text)"
            " VALUES (?,?,?,?,?,?)",
            (title, genre, characters, content_type, language, text)
        )
        conn.commit()
        return cur.lastrowid

def get_all_stories() -> list:
    """Return all saved stories as a list of dicts (newest first)."""
    init_db()
    with _get_conn() as conn:
        rows = conn.execute(
            "SELECT id, title, genre, content_type, language, saved_at, text"
            " FROM stories ORDER BY id DESC"
        ).fetchall()
    return [dict(r) for r in rows]

def delete_story(story_id: int):
    """Delete a story by id."""
    init_db()
    with _get_conn() as conn:
        conn.execute("DELETE FROM stories WHERE id = ?", (story_id,))
        conn.commit()


# ─────────────────────────────────────────────────────────────────────────────
# Phase 5 — Story Analysis & Advanced AI
# ─────────────────────────────────────────────────────────────────────────────

def rate_story(text: str) -> str:
    """Ask LLaMA to rate the story on Creativity, Emotion, and Plot."""
    prompt = f"""You are a professional literary critic and story analyst.

Read the following story and rate it strictly on these three criteria:
1. Creativity (originality of concept, unexpected twists)
2. Emotion (how emotionally engaging and resonant it is)
3. Plot (coherence, pacing, and story structure)

Rate each out of 10 and give ONE sentence of justification per criterion.

Return ONLY this exact format (no extra text):
Creativity: X/10 — [justification]
Emotion: X/10 — [justification]
Plot: X/10 — [justification]
Overall: X/10 — [one overall summary sentence]

Story:
\"\"\"
{text[:1500]}
\"\"\"
"""
    return ask_llama(prompt, temperature=0.3, max_tokens=200)


def analyze_story_stats(text: str) -> dict:
    """Return basic statistics about the story text (no LLM needed)."""
    import re
    words = text.split()
    word_count = len(words)
    reading_time_min = round(word_count / 200, 1)   # avg 200 wpm
    sentences = re.split(r'[.!?]+', text)
    sentence_count = len([s for s in sentences if s.strip()])
    char_mentions = {}
    # simple frequency count for capitalized words (likely names)
    candidates = re.findall(r'\b[A-Z][a-z]{2,}\b', text)
    for w in candidates:
        char_mentions[w] = char_mentions.get(w, 0) + 1
    # top 5
    top_chars = sorted(char_mentions.items(), key=lambda x: -x[1])[:5]
    return {
        "word_count":      word_count,
        "reading_time":    reading_time_min,
        "sentence_count":  sentence_count,
        "top_mentions":    top_chars,
    }


def generate_plot(genre: str, characters: str, theme: str) -> str:
    """Ask LLaMA to generate a 3-act plot structure."""
    prompt = f"""You are a professional Hollywood screenwriter and story structure expert.

Create a detailed 3-Act plot outline for a story with the following parameters:
Genre: {genre}
Main Characters: {characters}
Central Theme / Premise: {theme}

Structure it exactly as:
ACT 1 — SETUP
[3-4 bullet points]

ACT 2 — CONFRONTATION
[4-5 bullet points]

ACT 3 — RESOLUTION
[3-4 bullet points]

CLIMAX:
[1 sentence]

Return ONLY the structure above — no extra commentary.
"""
    return ask_llama(prompt, temperature=0.7, max_tokens=500)
